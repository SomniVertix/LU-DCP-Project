from docx import Document
import os
from os import walk

filenames = []
for dirpath, dirname, filename in walk('Word Docx/'):
    filenames.extend(filename)
    break

os.chdir('Word Docx/')
for doc in filenames:
    document = Document(doc)
    print (doc)
    for para in document.paragraphs:
        for run in para.runs:
            if ((run.underline == True or run.bold == True) and 'hours)' in para.text):
                print (para.text)
            
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if ((run.underline == True or run.bold == True) and 'hours)' in para.text):
                        print(paragraph.text)
#print (full_text)

